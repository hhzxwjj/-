#!/usr/bin/env python3
"""Convert all .md files in docs/ to .docx files."""
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def set_cell_border(cell, **kwargs):
    """Helper to set table cell borders."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = docx.oxml.OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = docx.oxml.OxmlElement(tag)
                tcBorders.append(element)
            for key in ["val", "color", "sz", "space"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def add_formatted_run(paragraph, text):
    """Parse inline markdown (bold, italic, code) and add runs."""
    # Pattern: code `...`, bold **...**, italic *...*
    # Process in order: code, bold+italic combined, bold, italic
    pos = 0
    while pos < len(text):
        # Code `...`
        m = re.search(r'`([^`]+)`', text[pos:])
        # Bold **...**
        mb = re.search(r'\*\*([^*]+)\*\*', text[pos:])
        # Italic *...* (but not **)
        mi = re.search(r'(?<!\*)\*([^*]+)\*(?!\*)', text[pos:])

        candidates = []
        if m:
            candidates.append((m.start(), m.end(), 'code', m.group(1)))
        if mb:
            candidates.append((mb.start(), mb.end(), 'bold', mb.group(1)))
        if mi:
            candidates.append((mi.start(), mi.end(), 'italic', mi.group(1)))

        if not candidates:
            paragraph.add_run(text[pos:])
            break

        candidates.sort(key=lambda x: x[0])
        start, end, fmt, content = candidates[0]

        if start > 0:
            paragraph.add_run(text[pos:pos + start])

        run = paragraph.add_run(content)
        if fmt == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xA9, 0x27, 0x21)
        elif fmt == 'bold':
            run.bold = True
        elif fmt == 'italic':
            run.italic = True

        pos += end


def md_to_docx(md_path, docx_path):
    """Convert a single markdown file to Word document."""
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_code():
        nonlocal code_lines, in_code_block
        if code_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('\n'.join(code_lines))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            code_lines = []
        in_code_block = False

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        # First row is header
        num_cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        table.style = 'Table Grid'
        for r_idx, row_cells in enumerate(table_rows):
            row = table.rows[r_idx]
            for c_idx, cell_text in enumerate(row_cells):
                cell = row.cells[c_idx]
                cell.text = cell_text.strip()
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                        run.font.name = '宋体'
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    if r_idx == 0:
                        for run in paragraph.runs:
                            run.bold = True
        doc.add_paragraph()
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                flush_code()
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty line
        if not line.strip():
            if in_table:
                flush_table()
            i += 1
            continue

        # Table separator line |---|---|
        if re.match(r'^\s*\|?\s*[:-]+\s*\|', line):
            i += 1
            continue

        # Table row
        if line.strip().startswith('|'):
            in_table = True
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            flush_table()

        # Headers
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            p = doc.add_heading(text, level=level)
            for run in p.runs:
                run.font.name = '黑体'
                run.font.name = 'Times New Roman'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^\s*([-*_]){3,}\s*$', line):
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Blockquote
        if line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            i += 1
            continue

        # List item
        list_match = re.match(r'^(\s*)((?:[-*+]|\d+\.)\s+)(.*)', line)
        if list_match:
            indent = len(list_match.group(1))
            bullet = list_match.group(2).strip()
            text = list_match.group(3)
            p = doc.add_paragraph(style='List Bullet' if bullet in ('-', '*', '+') else 'List Number')
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.15)
            add_formatted_run(p, text)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        add_formatted_run(p, line)
        i += 1

    if in_code_block:
        flush_code()
    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Converted: {md_path} -> {docx_path}")


def main():
    base_dir = Path('docs')
    md_files = sorted(base_dir.rglob('*.md'))

    for md_path in md_files:
        # Skip README.md in diagrams/plantuml
        if 'README' in md_path.name:
            continue

        # Determine output path: same relative path, but .docx extension
        rel = md_path.relative_to(base_dir)
        docx_name = md_path.stem + '.docx'
        docx_dir = md_path.parent
        docx_path = docx_dir / docx_name

        md_to_docx(str(md_path), str(docx_path))

    print(f"\nDone! Converted {len([f for f in md_files if 'README' not in f.name])} markdown files.")


if __name__ == '__main__':
    main()
